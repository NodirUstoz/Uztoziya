from rest_framework import generics, status, permissions
from rest_framework.decorators import api_view, permission_classes
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from django.shortcuts import get_object_or_404
from django.db.models import Q, Avg
from django.utils import timezone
import logging

logger = logging.getLogger(__name__)

from .models import Test, Question, Answer, TestAttempt, StudentAnswer, TestCategory
from .serializers import (
    TestCategorySerializer,
    TestSerializer,
    QuestionSerializer,
    AnswerSerializer,
    TestAttemptSerializer,
    StudentAnswerSerializer
)
from .ai_service import AITestGenerationService


class TestCategoryListView(generics.ListAPIView):
    """Test kategoriyalari ro'yxati"""
    queryset = TestCategory.objects.all()
    serializer_class = TestCategorySerializer
    permission_classes = [permissions.AllowAny]


class TestListView(generics.ListCreateAPIView):
    """Testlar ro'yxati va yaratish"""
    serializer_class = TestSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        queryset = Test.objects.filter(is_public=True, is_active=True)
        
        # Filtrlash
        category = self.request.query_params.get('category')
        subject = self.request.query_params.get('subject')
        grade_level = self.request.query_params.get('grade')
        difficulty = self.request.query_params.get('difficulty')
        search = self.request.query_params.get('search')
        
        if category:
            queryset = queryset.filter(category_id=category)
        if subject:
            queryset = queryset.filter(subject=subject)
        if grade_level:
            queryset = queryset.filter(grade_level=grade_level)
        if difficulty:
            queryset = queryset.filter(difficulty=difficulty)
        if search:
            queryset = queryset.filter(
                Q(title__icontains=search) |
                Q(description__icontains=search)
            )
        
        return queryset.order_by('-created_at')
    
    def perform_create(self, serializer):
        serializer.save(author=self.request.user)


class TestDetailView(generics.RetrieveUpdateDestroyAPIView):
    """Test tafsilotlari"""
    serializer_class = TestSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        return Test.objects.filter(
            Q(is_public=True) | Q(author=self.request.user)
        )


class TestCreateView(generics.CreateAPIView):
    """Yangi test yaratish"""
    serializer_class = TestSerializer
    permission_classes = [IsAuthenticated]
    
    def perform_create(self, serializer):
        serializer.save(author=self.request.user)
    
    def create(self, request, *args, **kwargs):
        """Test yaratish - AI funksiyasi bilan"""
        # AI test generation parametrlarini tekshirish
        ai_request = request.data.get('ai_generation')
        
        if ai_request:
            # AI yordamida test yaratish
            try:
                subject = request.data.get('subject')
                grade_level = request.data.get('grade_level')
                difficulty = request.data.get('difficulty', 'medium')
                num_questions = int(request.data.get('num_questions', 5))
                category_id = request.data.get('category_id')
                topic = request.data.get('topic', '')
                
                if not all([subject, grade_level]):
                    return Response({
                        'error': 'AI test yaratish uchun subject va grade_level majburiy'
                    }, status=status.HTTP_400_BAD_REQUEST)
                
                # AI service'ni ishga tushirish
                ai_service = AITestGenerationService()
                
                # AI yordamida savollar yaratish
                questions_data = ai_service.generate_test_questions(
                    subject=subject,
                    grade_level=grade_level,
                    difficulty=difficulty,
                    num_questions=num_questions
                )
                
                if not questions_data:
                    return Response({
                        'error': 'AI savollar yarata olmadi'
                    }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
                # Test sarlavhasi va tavsifini yaratish
                test_title = ai_service.generate_test_title(subject, grade_level, difficulty, topic)
                test_description = ai_service.generate_test_description(subject, grade_level, difficulty, len(questions_data))
                
                # Test yaratish
                test = Test.objects.create(
                    title=test_title,
                    description=test_description,
                    category_id=category_id,
                    subject=subject,
                    grade_level=grade_level,
                    difficulty=difficulty,
                    time_limit=len(questions_data) * 2,  # Har bir savol uchun 2 daqiqa
                    author=request.user
                )
                
                # Savollar va javoblarni yaratish
                for i, question_data in enumerate(questions_data):
                    question = Question.objects.create(
                        test=test,
                        question_text=question_data.get('question_text', ''),
                        question_type=question_data.get('question_type', 'single_choice'),
                        points=question_data.get('points', 1),
                        order=i + 1,
                        explanation=question_data.get('explanation', '')
                    )
                    
                    # Javoblarni yaratish
                    for j, answer_data in enumerate(question_data.get('answers', [])):
                        Answer.objects.create(
                            question=question,
                            answer_text=answer_data.get('answer_text', ''),
                            is_correct=answer_data.get('is_correct', False),
                            order=j + 1
                        )
                
                # Test statistikasini yangilash
                test.total_questions = test.questions.count()
                test.total_points = sum(q.points for q in test.questions.all())
                test.save()
                
                return Response({
                    'message': 'AI yordamida test muvaffaqiyatli yaratildi',
                    'test': TestSerializer(test).data,
                    'questions_count': len(questions_data)
                })
                
            except Exception as e:
                logger.error(f"AI test generation xatoligi: {e}")
                return Response({
                    'error': 'AI test yaratishda xatolik yuz berdi'
                }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        # Oddiy test yaratish
        return super().create(request, *args, **kwargs)


class TestUpdateView(generics.UpdateAPIView):
    """Testni yangilash"""
    serializer_class = TestSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        return Test.objects.filter(author=self.request.user)


class TestDeleteView(generics.DestroyAPIView):
    """Testni o'chirish"""
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        return Test.objects.filter(author=self.request.user)


class QuestionListView(generics.ListCreateAPIView):
    """Test savollari ro'yxati"""
    serializer_class = QuestionSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        test_id = self.kwargs['pk']
        test = get_object_or_404(Test, pk=test_id, author=self.request.user)
        return Question.objects.filter(test=test).order_by('order')
    
    def perform_create(self, serializer):
        test_id = self.kwargs['pk']
        test = get_object_or_404(Test, pk=test_id, author=self.request.user)
        serializer.save(test=test)


class QuestionDetailView(generics.RetrieveUpdateDestroyAPIView):
    """Savol tafsilotlari"""
    serializer_class = QuestionSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        return Question.objects.filter(test__author=self.request.user)


class QuestionCreateView(generics.CreateAPIView):
    """Yangi savol yaratish"""
    serializer_class = QuestionSerializer
    permission_classes = [IsAuthenticated]
    
    def perform_create(self, serializer):
        test_id = self.kwargs['pk']
        test = get_object_or_404(Test, pk=test_id, author=self.request.user)
        serializer.save(test=test)


class QuestionUpdateView(generics.UpdateAPIView):
    """Savolni yangilash"""
    serializer_class = QuestionSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        return Question.objects.filter(test__author=self.request.user)


class QuestionDeleteView(generics.DestroyAPIView):
    """Savolni o'chirish"""
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        return Question.objects.filter(test__author=self.request.user)


class TestAttemptListView(generics.ListAPIView):
    """Test topshirishlar ro'yxati"""
    serializer_class = TestAttemptSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        test_id = self.kwargs['pk']
        test = get_object_or_404(Test, pk=test_id, author=self.request.user)
        return TestAttempt.objects.filter(test=test).order_by('-started_at')


class TestAttemptDetailView(generics.RetrieveAPIView):
    """Test topshirish tafsilotlari"""
    serializer_class = TestAttemptSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        return TestAttempt.objects.filter(test__author=self.request.user)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def start_test(request, pk):
    """Testni boshlash"""
    test = get_object_or_404(Test, pk=pk, is_public=True, is_active=True)
    
    student_name = request.data.get('student_name')
    student_class = request.data.get('student_class', '')
    
    if not student_name:
        return Response({
            'error': 'O\'quvchi ismi kiritilishi kerak'
        }, status=status.HTTP_400_BAD_REQUEST)
    
    # Yangi test topshirish yaratish
    attempt = TestAttempt.objects.create(
        test=test,
        student_name=student_name,
        student_class=student_class,
        ip_address=request.META.get('REMOTE_ADDR')
    )
    
    # Test savollarini olish
    questions = test.questions.all().order_by('order')
    questions_data = []
    
    for question in questions:
        answers = question.answers.all().order_by('order')
        questions_data.append({
            'id': question.id,
            'question_text': question.question_text,
            'question_type': question.question_type,
            'points': question.points,
            'order': question.order,
            'image': question.image.url if question.image else None,
            'answers': [
                {
                    'id': answer.id,
                    'answer_text': answer.answer_text,
                    'order': answer.order
                }
                for answer in answers
            ]
        })
    
    return Response({
        'message': 'Test muvaffaqiyatli boshlandi',
        'attempt': TestAttemptSerializer(attempt).data,
        'questions': questions_data,
        'time_limit': test.time_limit
    })


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def submit_test(request, pk):
    """Testni topshirish"""
    test = get_object_or_404(Test, pk=pk, is_public=True, is_active=True)
    
    attempt_id = request.data.get('attempt_id')
    student_answers = request.data.get('student_answers', [])
    
    if not attempt_id:
        return Response({
            'error': 'Test topshirish ID si kiritilishi kerak'
        }, status=status.HTTP_400_BAD_REQUEST)
    
    try:
        attempt = TestAttempt.objects.get(id=attempt_id, test=test)
    except TestAttempt.DoesNotExist:
        return Response({
            'error': 'Test topshirish topilmadi'
        }, status=status.HTTP_404_NOT_FOUND)
    
    # Testni tugatish
    attempt.completed_at = timezone.now()
    attempt.is_completed = True
    
    # Javoblarni saqlash va baholash
    total_score = 0
    correct_answers = 0
    wrong_answers = 0
    
    for answer_data in student_answers:
        question_id = answer_data.get('question_id')
        selected_answers = answer_data.get('selected_answers', [])
        text_answer = answer_data.get('text_answer', '')
        
        try:
            question = Question.objects.get(id=question_id, test=test)
        except Question.DoesNotExist:
            continue
        
        # O'quvchi javobini saqlash
        student_answer = StudentAnswer.objects.create(
            attempt=attempt,
            question=question,
            text_answer=text_answer
        )
        
        # Tanlangan javoblarni saqlash
        for answer_id in selected_answers:
            try:
                answer = Answer.objects.get(id=answer_id, question=question)
                student_answer.selected_answers.add(answer)
            except Answer.DoesNotExist:
                continue
        
        # Javobni tekshirish
        is_correct = True
        correct_answers_list = question.answers.filter(is_correct=True)
        
        if question.question_type in ['single_choice', 'multiple_choice']:
            selected_correct = student_answer.selected_answers.filter(is_correct=True)
            if selected_correct.count() != correct_answers_list.count():
                is_correct = False
            elif not all(answer in selected_correct for answer in correct_answers_list):
                is_correct = False
        
        student_answer.is_correct = is_correct
        student_answer.points_earned = question.points if is_correct else 0
        student_answer.save()
        
        if is_correct:
            correct_answers += 1
            total_score += question.points
        else:
            wrong_answers += 1
    
    # Natijalarni saqlash
    attempt.score = total_score
    attempt.percentage = (total_score / test.total_points * 100) if test.total_points > 0 else 0
    attempt.save()
    
    return Response({
        'message': 'Test muvaffaqiyatli topshirildi',
        'attempt': TestAttemptSerializer(attempt).data,
        'results': {
            'score': total_score,
            'percentage': attempt.percentage,
            'correct_answers': correct_answers,
            'wrong_answers': wrong_answers,
            'total_questions': test.total_questions
        }
    })


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def search_tests(request):
    """Testlarni qidirish"""
    query = request.query_params.get('q', '')
    category = request.query_params.get('category')
    subject = request.query_params.get('subject')
    grade_level = request.query_params.get('grade')
    difficulty = request.query_params.get('difficulty')
    
    queryset = Test.objects.filter(is_public=True, is_active=True)
    
    if query:
        queryset = queryset.filter(
            Q(title__icontains=query) |
            Q(description__icontains=query)
        )
    
    if category:
        queryset = queryset.filter(category_id=category)
    if subject:
        queryset = queryset.filter(subject=subject)
    if grade_level:
        queryset = queryset.filter(grade_level=grade_level)
    if difficulty:
        queryset = queryset.filter(difficulty=difficulty)
    
    # Natijalarni tartiblash
    sort_by = request.query_params.get('sort', '-created_at')
    queryset = queryset.order_by(sort_by)
    
    serializer = TestSerializer(queryset, many=True)
    return Response(serializer.data)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def my_tests(request):
    """Foydalanuvchining testlari"""
    tests = Test.objects.filter(author=request.user).order_by('-created_at')
    serializer = TestSerializer(tests, many=True)
    return Response(serializer.data)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def test_stats(request, pk):
    """Test statistikasi"""
    test = get_object_or_404(Test, pk=pk, author=request.user)
    
    attempts = test.attempts.all()
    total_attempts = attempts.count()
    completed_attempts = attempts.filter(is_completed=True).count()
    avg_score = attempts.filter(is_completed=True).aggregate(
        avg_score=Avg('score')
    )['avg_score'] or 0
    avg_percentage = attempts.filter(is_completed=True).aggregate(
        avg_percentage=Avg('percentage')
    )['avg_percentage'] or 0
    
    return Response({
        'test': TestSerializer(test).data,
        'total_attempts': total_attempts,
        'completed_attempts': completed_attempts,
        'avg_score': round(avg_score, 2),
        'avg_percentage': round(avg_percentage, 2)
    })


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_ai_test(request):
    """AI yordamida test yaratish"""
    try:
        # Ma'lumotlarni olish
        subject = request.data.get('subject')
        grade_level = request.data.get('grade_level')
        difficulty = request.data.get('difficulty', 'medium')
        num_questions = int(request.data.get('num_questions', 5))
        category_id = request.data.get('category_id')
        topic = request.data.get('topic', '')
        
        if not all([subject, grade_level]):
            return Response({
                'error': 'Subject va grade_level majburiy'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # AI service'ni ishga tushirish
        ai_service = AITestGenerationService()
        
        # AI yordamida savollar yaratish
        questions_data = ai_service.generate_test_questions(
            subject=subject,
            grade_level=grade_level,
            difficulty=difficulty,
            num_questions=num_questions
        )
        
        if not questions_data:
            return Response({
                'error': 'AI savollar yarata olmadi'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        # Test sarlavhasi va tavsifini yaratish
        test_title = ai_service.generate_test_title(subject, grade_level, difficulty, topic)
        test_description = ai_service.generate_test_description(subject, grade_level, difficulty, len(questions_data))
        
        # Test yaratish
        test = Test.objects.create(
            title=test_title,
            description=test_description,
            category_id=category_id,
            subject=subject,
            grade_level=grade_level,
            difficulty=difficulty,
            time_limit=len(questions_data) * 2,  # Har bir savol uchun 2 daqiqa
            author=request.user
        )
        
        # Savollar va javoblarni yaratish
        for i, question_data in enumerate(questions_data):
            question = Question.objects.create(
                test=test,
                question_text=question_data.get('question_text', ''),
                question_type=question_data.get('question_type', 'single_choice'),
                points=question_data.get('points', 1),
                order=i + 1,
                explanation=question_data.get('explanation', '')
            )
            
            # Javoblarni yaratish
            for j, answer_data in enumerate(question_data.get('answers', [])):
                Answer.objects.create(
                    question=question,
                    answer_text=answer_data.get('answer_text', ''),
                    is_correct=answer_data.get('is_correct', False),
                    order=j + 1
                )
        
        # Test statistikasini yangilash
        test.total_questions = test.questions.count()
        test.total_points = sum(q.points for q in test.questions.all())
        test.save()
        
        return Response({
            'message': 'AI yordamida test muvaffaqiyatli yaratildi',
            'test': TestSerializer(test).data,
            'questions_count': len(questions_data)
        })
        
    except Exception as e:
        logger.error(f"AI test generation xatoligi: {e}")
        return Response({
            'error': 'AI test yaratishda xatolik yuz berdi'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def export_test_to_word(request):
    """Testni Word formatida export qilish"""
    try:
        from docx import Document
        from docx.shared import Inches
        from django.http import HttpResponse
        import io
        import logging
        
        logger = logging.getLogger(__name__)
        
        test_id = request.data.get('test_id')
        if not test_id:
            return Response({'error': 'Test ID kerak'}, status=status.HTTP_400_BAD_REQUEST)
        
        # Testni olish
        test = Test.objects.get(id=test_id, author=request.user)
        questions = test.questions.all().order_by('order')
        
        # Word hujjatini yaratish
        doc = Document()
        
        # Sarlavha
        title = doc.add_heading(test.title, 0)
        title.alignment = 1  # Markazga tekislash
        
        # Subject mapping
        subject_choices = {
            'mathematics': 'Matematika',
            'physics': 'Fizika',
            'chemistry': 'Kimyo',
            'biology': 'Biologiya',
            'geography': 'Geografiya',
            'history': 'Tarix',
            'literature': 'Adabiyot',
            'language': 'Til va adabiyot',
            'english': 'Ingliz tili',
            'russian': 'Rus tili',
            'computer_science': 'Informatika',
            'art': 'San\'at',
            'physical_education': 'Jismoniy tarbiya',
            'other': 'Boshqa'
        }
        
        # Difficulty mapping
        difficulty_choices = {
            'easy': 'Oson',
            'medium': 'O\'rta',
            'hard': 'Qiyin'
        }
        
        # Test ma'lumotlari
        doc.add_paragraph(f"Fan: {subject_choices.get(test.subject, test.subject)}")
        doc.add_paragraph(f"Sinf: {test.grade_level}")
        doc.add_paragraph(f"Qiyinlik: {difficulty_choices.get(test.difficulty, test.difficulty)}")
        doc.add_paragraph(f"Vaqt chegarasi: {test.time_limit} daqiqa")
        doc.add_paragraph(f"Savollar soni: {questions.count()}")
        
        if test.description:
            doc.add_paragraph(f"Tavsif: {test.description}")
        
        doc.add_paragraph()  # Bo'sh qator
        
        # Savollar
        for i, question in enumerate(questions, 1):
            # Savol matni
            doc.add_heading(f"Savol {i}", level=2)
            doc.add_paragraph(question.question_text)
            
            # Javoblar
            answers = question.answers.all().order_by('order')
            for j, answer in enumerate(answers):
                letter = chr(65 + j)  # A, B, C, D
                doc.add_paragraph(f"{letter}) {answer.answer_text}")
            
            # Tushuntirish
            if question.explanation:
                doc.add_paragraph(f"Tushuntirish: {question.explanation}")
            
            doc.add_paragraph()  # Bo'sh qator
        
        # Javoblar jadvali
        doc.add_heading("Javoblar jadvali", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Jadval sarlavhalari
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Savol'
        hdr_cells[1].text = 'To\'g\'ri javob'
        
        # Javoblar
        for i, question in enumerate(questions, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            
            # To'g'ri javobni topish
            correct_answer = question.answers.filter(is_correct=True).first()
            if correct_answer:
                answer_order = correct_answer.order
                row_cells[1].text = chr(64 + answer_order)  # A, B, C, D
            else:
                row_cells[1].text = "Javob yo'q"
        
        # Response yaratish
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{test.title}.docx"'
        
        # Hujjatni saqlash
        doc.save(response)
        
        return response
        
    except Test.DoesNotExist:
        return Response({'error': 'Test topilmadi'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        logger.error(f"Word export xatoligi: {e}")
        return Response({'error': 'Word export xatoligi'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)